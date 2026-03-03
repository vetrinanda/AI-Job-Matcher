"""Resume parser — extracts text from PDF and DOCX files."""

import io
from PyPDF2 import PdfReader
from docx import Document


def parse_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF file."""
    reader = PdfReader(io.BytesIO(file_bytes))
    text_parts = []
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text_parts.append(page_text)
    return "\n".join(text_parts)


def parse_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file."""
    doc = Document(io.BytesIO(file_bytes))
    text_parts = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)
    return "\n".join(text_parts)


def parse_resume(file_bytes: bytes, filename: str) -> str:
    """Parse a resume file and return extracted text.

    Args:
        file_bytes: Raw file content
        filename: Original filename (used to determine format)

    Returns:
        Extracted text from the resume

    Raises:
        ValueError: If the file format is not supported
    """
    lower_name = filename.lower()
    if lower_name.endswith(".pdf"):
        return parse_pdf(file_bytes)
    elif lower_name.endswith(".docx"):
        return parse_docx(file_bytes)
    elif lower_name.endswith(".doc"):
        raise ValueError("Legacy .doc format is not supported. Please convert to .docx or PDF.")
    else:
        raise ValueError(f"Unsupported file format: {filename}. Please upload a PDF or DOCX file.")
